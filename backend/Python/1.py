from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

def set_shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def make_table(uc):
    t = doc.add_table(rows=0, cols=2)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    def row(label, val, merge=False, italic=False):
        r = t.add_row()
        if merge:
            m = r.cells[0].merge(r.cells[1])
            m.text = ''
            run = m.paragraphs[0].add_run(label)
            run.bold = True; run.italic = italic
            run.font.name='Times New Roman'; run.font.size=Pt(11)
            set_shade(m, 'D9E1F2')
        else:
            r.cells[0].text=''
            run1 = r.cells[0].paragraphs[0].add_run(label)
            run1.bold=True; run1.font.name='Times New Roman'; run1.font.size=Pt(11)
            set_shade(r.cells[0], 'D9E1F2')
            r.cells[1].text=''
            run2 = r.cells[1].paragraphs[0].add_run(val)
            run2.font.name='Times New Roman'; run2.font.size=Pt(11)

    def merged(text, bold=False, italic=False):
        r = t.add_row()
        m = r.cells[0].merge(r.cells[1])
        m.text=''
        run = m.paragraphs[0].add_run(text)
        run.bold=bold; run.italic=italic
        run.font.name='Times New Roman'; run.font.size=Pt(11)

    def post(text):
        r = t.add_row()
        m = r.cells[0].merge(r.cells[1])
        m.text=''
        p = m.paragraphs[0]
        r1 = p.add_run('Post Condition: '); r1.bold=True
        r1.font.name='Times New Roman'; r1.font.size=Pt(11)
        r2 = p.add_run(text)
        r2.font.name='Times New Roman'; r2.font.size=Pt(11)

    row('Use Case Name', uc['name'])
    row('Description', uc['description'])
    row('Primary Actor', uc['primary_actor'])
    row('Secondary Actor', uc['secondary_actor'])
    row('Precondition', uc['precondition'])
    row('Dependency', uc['dependency'])
    row('Generalization', uc['generalization'])
    row('Basic Flow', '', merge=True)
    merged('Main Successful Path:', bold=True)
    for i,s in enumerate(uc['basic_flow'],1): merged(f'{i}. {s}')
    post(uc['post_condition_basic'])
    row('Alternative Flow', '', merge=True)
    merged(uc['alt_flow_intro'], italic=True)
    for i,s in enumerate(uc['alternative_flow'],1): merged(f'{i}. {s}')
    post(uc['post_condition_alt'])

use_cases = [
    {'name':'Authentication & Access Management','description':'Handles user registration, login, and logout to provide secure access to the system.','primary_actor':'User','secondary_actor':'System, Admin','precondition':'User must have valid credentials or a registered account.','dependency':'Includes password reset and account verification','generalization':'None','basic_flow':['User opens the login or registration page.','User enters credentials or uses Google sign-in.','System verifies user identity.','System grants access and redirects to dashboard.','User can logout when required.'],'post_condition_basic':'User is successfully authenticated and gains access to authorized features.','alt_flow_intro':'Specifically applies when login fails or password is forgotten.','alternative_flow':['User selects "Forgot Password" option.','System requests email verification.','User resets password through provided link.','User logs in again using new credentials.'],'post_condition_alt':'User regains account access after successful password reset.'},
    {'name':'User Profile & Medical Information Management','description':'Allows users to view and update personal details and medical history.','primary_actor':'User','secondary_actor':'Admin','precondition':'User must be logged in.','dependency':'Includes data validation and consent confirmation','generalization':'None','basic_flow':['User opens profile section.','System displays existing information.','User updates medical or personal details.','System validates and saves updated data.','Confirmation message is shown.'],'post_condition_basic':'User profile information is updated and stored securely.','alt_flow_intro':'Applies when user cancels or enters invalid data.','alternative_flow':['User starts editing profile.','User cancels changes or enters invalid input.','System rejects invalid data or discards changes.','User returns to profile view.'],'post_condition_alt':'No changes are saved; previous data remains unchanged.'},
    {'name':'Symptom Submission & Onboarding Assessment','description':'Collects user symptoms through structured questionnaires for analysis.','primary_actor':'User','secondary_actor':'System','precondition':'User must be logged in or in onboarding phase.','dependency':'Includes questionnaire module','generalization':'None','basic_flow':['System presents symptom questionnaire.','User fills in required symptom details.','User submits responses.','System stores submitted data.','System confirms successful submission.'],'post_condition_basic':'Symptom data is recorded for analysis and recommendations.','alt_flow_intro':'Applies when questionnaire is incomplete.','alternative_flow':['User skips or misses required fields.','System prompts for missing information.','User completes remaining questions.','System accepts final submission.'],'post_condition_alt':'Complete symptom data is successfully stored.'},
    {'name':'Diabetes Risk Assessment','description':'Evaluates user data to generate diabetes risk level using AI model.','primary_actor':'User','secondary_actor':'System','precondition':'User must have submitted symptom and profile data.','dependency':'Includes ML prediction engine','generalization':'None','basic_flow':['User requests risk assessment or completes questionnaire.','System retrieves user medical and symptom data.','ML model processes data.','System generates risk score and explanation.','Results are displayed to user.'],'post_condition_basic':'User receives diabetes risk result and stored analysis report.','alt_flow_intro':'Applies when system cannot generate prediction.','alternative_flow':['System fails to compute risk due to missing or insufficient data.','System requests additional information.','User provides required data.','System reprocesses and generates result.'],'post_condition_alt':'Risk assessment is generated after completing required data.'},
    {'name':'Health Metrics Logging & Tracking','description':'Allows users to record and track health indicators over time.','primary_actor':'User','secondary_actor':'System / Device (optional)','precondition':'User must be authenticated.','dependency':'Includes data validation and trend analysis','generalization':'None','basic_flow':['User opens health metrics section.','User enters or syncs health data.','System validates and stores data.','System updates health trends.','User views updated history.'],'post_condition_basic':'Health metrics are saved and available for tracking.','alt_flow_intro':'Applies when invalid data is entered.','alternative_flow':['User enters incorrect or incomplete data.','System detects validation error.','System requests correction.','User resubmits valid data.'],'post_condition_alt':'Only valid health metrics are stored.'},
    {'name':'AI Chat Assistant (RAG-based Support)','description':'User interacts with an AI chat system to ask health-related questions and receive personalized responses based on their profile and system knowledge.','primary_actor':'User','secondary_actor':'System (AI engine, knowledge base)','precondition':'User must be authenticated and have an active session.','dependency':'Includes user profile context and knowledge retrieval system','generalization':'None','basic_flow':['User opens the AI chat interface.','User enters a health-related question.','System retrieves relevant user data and knowledge base information.','System generates a personalized response.','Response is shown to the user in chat.'],'post_condition_basic':'User receives an AI-generated personalized response.','alt_flow_intro':'Specifically applies when system cannot generate confident response.','alternative_flow':['User enters a question.','System fails to find strong relevant data.','System provides general guidance instead of personalized answer.','System suggests consulting a healthcare professional.'],'post_condition_alt':'User receives a fallback response with general advice.'},
    {'name':'Personalized Diet Plan Generation (Monthly)','description':'System generates a personalized monthly diet plan based on user health data, preferences, and medical conditions.','primary_actor':'User','secondary_actor':'System (Recommendation Engine)','precondition':'User profile and dietary preferences must be available.','dependency':'Uses user health data and nutrition rules engine','generalization':'None','basic_flow':['User requests a diet plan.','System collects user medical and preference data.','System generates a full monthly diet plan.','System displays the plan to the user.','User downloads or saves the plan.'],'post_condition_basic':'Monthly diet plan is generated and stored.','alt_flow_intro':'Specifically applies when conflicting dietary conditions exist.','alternative_flow':['User requests diet plan.','System detects conflicting dietary restrictions.','System asks user for clarification or preference adjustment.','System regenerates updated diet plan.'],'post_condition_alt':'Adjusted monthly diet plan is generated.'},
    {'name':'Exercise Plan Generation','description':'System creates a personalized exercise plan based on user fitness level, medical condition, and goals.','primary_actor':'User','secondary_actor':'System (Recommendation Engine)','precondition':'User health and fitness profile must exist.','dependency':'Uses medical constraints and activity rules','generalization':'None','basic_flow':['User requests an exercise plan.','System retrieves user fitness and medical data.','System generates personalized exercise recommendations.','System displays the plan to the user.','User saves or accepts the plan.'],'post_condition_basic':'Exercise plan is created and stored.','alt_flow_intro':'Specifically applies when medical restriction exists.','alternative_flow':['User requests exercise plan.','System detects physical limitation or injury.','System adjusts exercise intensity or type.','System regenerates safe exercise plan.'],'post_condition_alt':'Safe exercise plan is generated based on constraints.'},
    {'name':'Lifestyle Tips & Habit Suggestions','description':'System suggests personalized lifestyle improvements such as sleep, hydration, and daily habits.','primary_actor':'User','secondary_actor':'System (Suggestion Engine)','precondition':'User must have profile and basic health data.','dependency':'Uses historical data and symptom trends','generalization':'None','basic_flow':['User opens lifestyle suggestions section.','System analyzes user health and behavior data.','System generates personalized habit suggestions.','System displays recommendations to the user.','User selects and saves preferred habits.'],'post_condition_basic':'Selected habits are stored for tracking.','alt_flow_intro':'Specifically applies when user rejects suggestions.','alternative_flow':['User views suggestions.','User declines or ignores recommendations.','System records preference.','System updates future suggestion logic.'],'post_condition_alt':'System adapts future suggestions based on user behavior.'},
    {'name':'Weekly Health Priorities System','description':'System generates weekly health priorities based on user plans, goals, and recent health data.','primary_actor':'User','secondary_actor':'System (Planning Engine)','precondition':'User must have active plans or recent health data.','dependency':'Integrates diet, exercise, and symptom tracking modules','generalization':'None','basic_flow':['System generates weekly health priorities.','User reviews suggested priorities.','User accepts or modifies priorities.','System saves weekly plan.','System tracks progress throughout the week.'],'post_condition_basic':'Weekly priorities are active and monitored.','alt_flow_intro':'Specifically applies when user requests change.','alternative_flow':['User reviews weekly priorities.','User requests different focus or changes.','System regenerates updated priorities.','System saves revised plan.'],'post_condition_alt':'Updated weekly priorities are stored.'},
    {'name':'Reports & History Management (PDF Download & Export)','description':'User views past health records, assessments, and generated plans, and can download or export them as reports.','primary_actor':'User','secondary_actor':'System (Report Generator)','precondition':'User must be authenticated and have stored historical data.','dependency':'Includes report generation and data aggregation modules','generalization':'None','basic_flow':['User opens the reports/history section.','User selects data type and time range.','System compiles relevant health data.','System generates a structured report.','User downloads or views the report.'],'post_condition_basic':'Report is generated and accessible to the user.','alt_flow_intro':'Specifically applies when no data is available.','alternative_flow':['User requests report.','System checks for available data.','System finds no records for selected range.','System notifies user and suggests changing filters.'],'post_condition_alt':'No report is generated due to lack of data.'},
    {'name':'User Profile & Medical Information Management','description':'User manages personal details, medical history, allergies, and related health information.','primary_actor':'User','secondary_actor':'Admin (support access if required)','precondition':'User must be logged in.','dependency':'Includes data validation and secure storage','generalization':'None','basic_flow':['User opens profile settings.','User views existing personal and medical data.','User updates required information.','System validates input data.','System saves updated profile.'],'post_condition_basic':'Profile data is updated successfully.','alt_flow_intro':'Specifically applies when invalid data is entered.','alternative_flow':['User enters profile data.','System detects invalid or incomplete fields.','System prompts user to correct information.','User updates and resubmits data.'],'post_condition_alt':'Only valid data is saved in the system.'},
    {'name':'Health Metrics Logging & Tracking','description':'User records health metrics such as glucose levels, weight, and blood pressure, and tracks progress over time.','primary_actor':'User','secondary_actor':'External devices (optional integration)','precondition':'User must be authenticated.','dependency':'Includes data storage and trend analysis module','generalization':'None','basic_flow':['User opens metrics logging section.','User enters health metric values or syncs device data.','System validates and stores data.','System updates health trends.','User views updated charts and progress.'],'post_condition_basic':'Health metrics are stored and visualized.','alt_flow_intro':'Specifically applies when invalid input is entered.','alternative_flow':['User enters metric values.','System detects invalid format or range.','System requests correction.','User updates and resubmits data.'],'post_condition_alt':'Only valid metrics are stored.'},
    {'name':'Audit Logging & System Monitoring','description':'System records and monitors user and system activities for security, compliance, and debugging purposes.','primary_actor':'Admin / Super Admin','secondary_actor':'System Monitoring Services','precondition':'Admin must have monitoring privileges.','dependency':'Includes logging and audit trail system','generalization':'None','basic_flow':['Admin opens audit log dashboard.','Admin selects filters (time, user, action type).','System displays relevant logs.','Admin reviews system activities.','Admin exports or flags important logs.'],'post_condition_basic':'Audit logs are reviewed and recorded.','alt_flow_intro':'Specifically applies when anomaly is detected.','alternative_flow':['System detects unusual activity pattern.','System highlights anomaly in logs.','Admin investigates flagged activity.','System records investigation action.'],'post_condition_alt':'Suspicious activity is flagged and documented.'}
]

t = doc.add_paragraph()
r = t.add_run('3.2 Use Case Description'); r.bold=True; r.font.size=Pt(16); r.font.name='Times New Roman'
t = doc.add_paragraph()
r = t.add_run('Table 3.1: Textual Description of Use Cases'); r.bold=True; r.font.size=Pt(13); r.font.name='Times New Roman'

for i, uc in enumerate(use_cases, 1):
    h = doc.add_paragraph()
    r = h.add_run(f'Use Case {i}: {uc["name"]}')
    r.bold=True; r.font.size=Pt(13); r.font.name='Times New Roman'
    make_table(uc)
    doc.add_paragraph()

doc.save('Use_Cases_Document.docx')
print("Done! File saved as Use_Cases_Document.docx")